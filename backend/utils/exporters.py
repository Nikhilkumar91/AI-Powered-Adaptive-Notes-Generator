from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.config import DOCX_DIR, PDF_DIR

PDF_DIR.mkdir(parents=True, exist_ok=True)
DOCX_DIR.mkdir(parents=True, exist_ok=True)


def _stamp() -> str:
    return datetime.now().strftime('%Y%m%d_%H%M%S')


def create_pdf(topic: str, notes: dict, level: str) -> Path:
    path = PDF_DIR / f'notes_{level}_{_stamp()}.pdf'
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        textColor=colors.HexColor('#0f4c81'),
        fontSize=22,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], textColor=colors.HexColor('#17324d'))
    body_style = styles['BodyText']
    body_style.leading = 15
    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=0.7 * inch, rightMargin=0.7 * inch)
    story = [
        Paragraph(topic, title_style),
        Paragraph(f'Level: {level.title()}', body_style),
        Spacer(1, 12),
        Paragraph('Introduction', heading_style),
        Paragraph(notes.get('introduction', ''), body_style),
        Spacer(1, 8),
        Paragraph('Key Highlights', heading_style),
    ]
    for item in notes.get('highlights', []):
        story.append(Paragraph(f'- {item}', body_style))
    story.extend([Spacer(1, 8), Paragraph('Learning Structure', heading_style)])
    for item in notes.get('structure', []):
        story.append(Paragraph(f'- {item}', body_style))
    story.extend([Spacer(1, 8), Paragraph('Summary', heading_style), Paragraph(notes.get('summary', ''), body_style)])
    doc.build(story)
    return path


def create_docx(topic: str, notes: dict, level: str) -> Path:
    path = DOCX_DIR / f'notes_{level}_{_stamp()}.docx'
    document = Document()
    heading = document.add_heading(topic, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f'Level: {level.title()}')
    document.add_heading('Introduction', level=1)
    document.add_paragraph(notes.get('introduction', ''))
    document.add_heading('Key Highlights', level=1)
    for item in notes.get('highlights', []):
        document.add_paragraph(item, style='List Bullet')
    document.add_heading('Learning Structure', level=1)
    for item in notes.get('structure', []):
        document.add_paragraph(item, style='List Number')
    document.add_heading('Summary', level=1)
    document.add_paragraph(notes.get('summary', ''))
    footer = document.add_paragraph(f'Generated on {datetime.now():%Y-%m-%d %H:%M:%S}')
    for run in footer.runs:
        run.font.size = Pt(9)
        run.italic = True
    document.save(path)
    return path
